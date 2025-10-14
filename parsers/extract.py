import re
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from PyPDF2 import PdfReader
import docx
import spacy
from spacy.matcher import Matcher
import phonenumbers
from datetime import datetime
import json
import subprocess
import fitz  

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ResumeParser:
    """Production-grade resume parser with NLP capabilities"""
    
    def __init__(self, spacy_model: str = "en_core_web_sm"):
        """
        Initialize the parser with spaCy model
        
        Args:
            spacy_model: spaCy model name (default: en_core_web_sm)
        """
        try:
            self.nlp = spacy.load(spacy_model)
        except OSError:
            logger.warning(f"Model {spacy_model} not found. Downloading...")
            subprocess.run(["python", "-m", "spacy", "download", spacy_model], check=True)
            self.nlp = spacy.load(spacy_model)
        
        self.matcher = Matcher(self.nlp.vocab)
        self._setup_matchers()
        self._load_skill_database()
        
    def _setup_matchers(self):
        """Setup spaCy matchers for pattern recognition"""
        # Education patterns
        education_patterns = [
            [{"LOWER": {"IN": ["bachelor", "master", "phd", "doctorate"]}}],
            [{"TEXT": {"REGEX": "B\\.?(Tech|E|Sc|A|Com)"}}],
            [{"TEXT": {"REGEX": "M\\.?(Tech|E|Sc|BA|Com|CA)"}}]
        ]
        self.matcher.add("EDUCATION", education_patterns)
        
    def _load_skill_database(self):
        """Load comprehensive skill database"""
        self.skills_db = {
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c',
                'ruby', 'php', 'swift', 'kotlin', 'go', 'rust', 'scala', 'r',
                'perl', 'dart', 'objective-c', 'sql', 'bash', 'powershell'
            ],
            'frameworks': [
                'react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'spring',
                'spring boot', 'express', 'node.js', 'next.js', 'nuxt.js',
                'laravel', 'symfony', 'rails', 'asp.net', '.net core', 'flutter',
                'react native', 'tensorflow', 'pytorch', 'keras', 'scikit-learn'
            ],
            'cloud_devops': [
                'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes',
                'jenkins', 'gitlab ci', 'github actions', 'terraform', 'ansible',
                'circleci', 'travis ci', 'heroku', 'digitalocean'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
                'cassandra', 'dynamodb', 'oracle', 'sql server', 'sqlite',
                'neo4j', 'couchdb', 'mariadb'
            ],
            'tools_technologies': [
                'git', 'jira', 'confluence', 'agile', 'scrum', 'kanban',
                'rest api', 'graphql', 'microservices', 'ci/cd', 'tdd',
                'machine learning', 'deep learning', 'nlp', 'computer vision',
                'data analysis', 'data science', 'big data', 'hadoop', 'spark',
                'tableau', 'power bi', 'excel', 'sap', 'salesforce'
            ]
        }
        
        # Flatten all skills
        self.all_skills = []
        for category in self.skills_db.values():
            self.all_skills.extend(category)
        
        # Create regex pattern for skills
        self.skills_pattern = re.compile(
            r'\b(' + '|'.join(re.escape(skill) for skill in self.all_skills) + r')\b',
            re.IGNORECASE
        )
    
    def read_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF primarily, fallback to PyPDF2."""
        text = ""

        # ---------- Attempt 1: PyMuPDF (best for modern PDFs) ----------
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text") or ""
                text += page_text + "\n"
            doc.close()
            if text.strip():
                logger.info(f"[INFO] Extracted {len(text)} characters via PyMuPDF")
                return text
        except Exception as e:
            logger.warning(f"[WARN] PyMuPDF extraction failed: {e}")

        # ---------- Attempt 2: PyPDF2 (fallback) ----------
        try:
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
            if text.strip():
                logger.info(f"[INFO] Extracted {len(text)} characters via PyPDF2 fallback")
                return text
        except Exception as e:
            logger.error(f"[ERROR] PyPDF2 extraction failed: {e}")

        logger.warning(f"[WARN] No text extracted from {file_path}")
        return text

    def read_docx(self, file_path: str) -> str:
        """Extract text from DOCX with error handling"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            raise
        return text
    
    def read_txt(self, file_path: str) -> str:
        """Read text file with encoding detection"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Unable to decode file with supported encodings")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text based on file extension"""
        path = Path(file_path)
        logger.info(f"[DEBUG] Reading file from: {file_path}")
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            logger.info(f"[DEBUG] Detected extension: {extension}")
            return self.read_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            logger.info(f"[DEBUG] Detected extension: {extension}")
            return self.read_docx(file_path)
        elif extension == '.txt':
            return self.read_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def extract_name(self, text: str) -> Optional[str]:
        """Extract name using NLP with improved accuracy"""
        doc = self.nlp(text[:1000])  # Check first 1000 chars
        
        # Look for PERSON entities in the first few lines
        lines = text.strip().split('\n')[:10]
        
        for line in lines:
            line = line.strip()
            if not line or len(line) > 100:
                continue
            
            # Skip lines with common resume headers or contact info
            if re.search(r'resume|curriculum|vitae|email|phone|address|linkedin|github', 
                        line, re.IGNORECASE):
                continue
            
            # Use spaCy NER
            line_doc = self.nlp(line)
            for ent in line_doc.ents:
                if ent.label_ == "PERSON":
                    return ent.text.strip()
            
            # Fallback: check if line looks like a name
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() and word.replace('.', '').isalpha() 
                      for word in words):
                    return line
        
        return None
    
    def extract_email(self, text: str) -> List[str]:
        """Extract all email addresses"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return list(set(emails))  # Remove duplicates
    
    def extract_phone(self, text: str) -> List[str]:
        """Extract phone numbers using phonenumbers library"""
        phones = []
        
        # Try to find phone numbers with country context
        try:
            for match in phonenumbers.PhoneNumberMatcher(text, "IN"):
                phones.append(phonenumbers.format_number(
                    match.number, 
                    phonenumbers.PhoneNumberFormat.INTERNATIONAL
                ))
        except Exception as e:
            logger.warning(f"Phone number extraction error: {e}")
        
        # Also use regex for backup
        phone_patterns = [
            r'\+?91[-.\s]?[6-9]\d{9}',  # Indian format
            r'\+?1?[-.]?\(?([0-9]{3})\)?[-.]?([0-9]{3})[-.]?([0-9]{4})',
            r'\b[6-9]\d{9}\b',  # Indian 10-digit
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    phone = ''.join(match)
                else:
                    phone = match
                phone = re.sub(r'[-.\s()]', '', phone)
                if len(phone) >= 10:
                    phones.append(phone)
        
        # Deduplicate and clean
        unique_phones = []
        seen = set()
        for phone in phones:
            clean = re.sub(r'\D', '', phone)
            if clean not in seen:
                seen.add(clean)
                unique_phones.append(phone)
        
        return unique_phones
    
    def extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract skills categorized by type"""
        text_lower = text.lower()
        
        categorized_skills = {}
        
        for category, skills_list in self.skills_db.items():
            found_skills = []
            for skill in skills_list:
                # Use word boundary matching for accuracy
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.append(skill)
            
            if found_skills:
                categorized_skills[category] = found_skills
        
        # Also extract all unique skills as a flat list
        all_found = []
        for skills in categorized_skills.values():
            all_found.extend(skills)
        
        return {
            'categorized': categorized_skills,
            'all_skills': list(set(all_found))
        }
    
    def extract_experience(self, text: str) -> Dict[str, any]:
        """Extract detailed experience information"""
        experience_info = {
            'total_years': None,
            'work_history': []
        }
        
        # Extract total years - more strict patterns
        exp_patterns = [
            r'(?:total\s+)?experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)',
            r'(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:total\s+)?experience',
        ]
        
        for pattern in exp_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                years = float(match.group(1))
                if years > 0 and years < 50:  # Sanity check
                    experience_info['total_years'] = years
                    break
        
        # Extract work history with dates - more specific
        date_pattern = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\s*[-–—]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current)'
        dates = re.findall(date_pattern, text, re.IGNORECASE)
        
        if dates and len(dates) > 0:
            experience_info['date_range'] = f"{dates[0][0]} - {dates[-1][1]}"
        
        return experience_info
    
    def extract_education(self, text: str) -> List[Dict]:
        """Extract education with multiple fallback strategies"""
        education_list = []
        seen_degrees = set()
        
        # Strategy 1: Look for EDUCATION section header
        education_section = ""
        edu_header_match = re.search(
            r'(?:EDUCATION|ACADEMIC|QUALIFICATION).*?(?=\n(?:EXPERIENCE|WORK|PROJECTS|SKILLS|CERTIFICATIONS|$))',
            text,
            re.IGNORECASE | re.DOTALL
        )
        if edu_header_match:
            education_section = edu_header_match.group(0)
            logger.info(f"Found education section: {len(education_section)} chars")
        else:
            education_section = text  # Fallback to full text
        
        # Strategy 2: Multiple degree patterns with variations
        degree_patterns = [
            # Verbose patterns
            (r'Bachelor\s+of\s+(?:Technology|Engineering|Science|Arts|Commerce)\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;)', 'Bachelor'),
            (r'Master\s+of\s+(?:Technology|Engineering|Science|Business Administration|Arts|Commerce)\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;)', 'Master'),
            
            # Abbreviated patterns
            (r'B\.?\s*Tech\.?\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;|\d{4})', 'B.Tech'),
            (r'B\.?\s*E\.?\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;|\d{4})', 'B.E'),
            (r'B\.?\s*Sc\.?\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;|\d{4})', 'B.Sc'),
            (r'M\.?\s*Tech\.?\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;|\d{4})', 'M.Tech'),
            (r'M\.?\s*E\.?\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;|\d{4})', 'M.E'),
            (r'M\.?\s*Sc\.?\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;|\d{4})', 'M.Sc'),
            (r'MBA\s*(?:in\s+)?([A-Za-z\s&,()-]+)?', 'MBA'),
            (r'MCA', 'MCA'),
            (r'BCA', 'BCA'),
            (r'Ph\.?D\.?\s+(?:in\s+)?([A-Za-z\s&,()-]+?)(?=\n|$|,|\||;)', 'PhD'),
        ]
        
        for pattern, degree_name in degree_patterns:
            matches = re.finditer(pattern, education_section, re.IGNORECASE)
            for match in matches:
                full_match = match.group(0).strip()
                
                # Skip if too short
                if len(full_match) < 3:
                    continue
                
                # Create a normalized key for deduplication
                normalized_key = full_match.lower().replace('.', '').replace(' ', '').replace('\n', '')
                if normalized_key in seen_degrees:
                    continue
                
                seen_degrees.add(normalized_key)
                
                education_entry = {
                    'degree': degree_name,
                    'field': None
                }
                
                # Extract field of study
                if match.lastindex and match.lastindex >= 1:
                    field = match.group(1)
                    if field:
                        field = field.strip().strip(',;|')
                        if len(field) > 3 and len(field) < 100:
                            education_entry['field'] = field
                
                # Find associated institution in nearby context (500 chars)
                start = max(0, match.start() - 100)
                end = min(len(education_section), match.end() + 400)
                context = education_section[start:end]
                
                # Institution patterns - look for university/college/institute
                institution_patterns = [
                    r'([A-Z][A-Za-z\s&,.-]+?(?:University|Institute|College|School)(?:\s+of\s+[A-Za-z\s]+)?)',
                    r'(IIT|NIT|IIIT|VIT|BITS|IIM|AIIMS)\s+([A-Za-z]+)',
                ]
                
                for inst_pattern in institution_patterns:
                    college_matches = re.finditer(inst_pattern, context, re.IGNORECASE)
                    for college_match in college_matches:
                        institution = college_match.group(0).strip()
                        institution = re.sub(r'\s+', ' ', institution)
                        institution = institution.rstrip(',;.\n|')
                        
                        if 5 < len(institution) < 150:
                            education_entry['institution'] = institution
                            break
                    if 'institution' in education_entry:
                        break
                
                # Extract graduation year
                year_pattern = r'\b(19\d{2}|20[0-2]\d)\b'
                years = re.findall(year_pattern, context)
                if years:
                    if len(years) == 1:
                        education_entry['year'] = years[0]
                    elif len(years) >= 2:
                        education_entry['year'] = f"{years[0]} - {years[-1]}"
                
                # Extract CGPA/GPA/Percentage
                grade_patterns = [
                    r'(?:CGPA|GPA)[:\s]*(\d+\.?\d*)\s*(?:/\s*(\d+\.?\d*))?',
                    r'(?:Percentage|%)[:\s]*(\d+\.?\d*)\s*%?',
                ]
                
                for grade_pattern in grade_patterns:
                    grade_match = re.search(grade_pattern, context, re.IGNORECASE)
                    if grade_match:
                        if grade_match.lastindex >= 2 and grade_match.group(2):
                            education_entry['grade'] = f"{grade_match.group(1)}/{grade_match.group(2)}"
                        else:
                            education_entry['grade'] = grade_match.group(1)
                        break
                
                education_list.append(education_entry)
        
        logger.info(f"Extracted {len(education_list)} education entries")
        return education_list
    
    def extract_work_experience(self, text: str) -> List[Dict]:
        """Extract work experience with multiple strategies"""
        work_history = []
        seen_entries = set()
        
        # Strategy 1: Look for EXPERIENCE/WORK section
        experience_section = ""
        exp_header_match = re.search(
            r'(?:EXPERIENCE|EMPLOYMENT|WORK HISTORY|PROFESSIONAL EXPERIENCE).*?(?=\n(?:EDUCATION|PROJECTS|SKILLS|CERTIFICATIONS|$))',
            text,
            re.IGNORECASE | re.DOTALL
        )
        if exp_header_match:
            experience_section = exp_header_match.group(0)
            logger.info(f"Found experience section: {len(experience_section)} chars")
        else:
            experience_section = text  # Fallback to full text
        
        # Strategy 2: Multiple patterns for job titles and companies
        
        # Pattern A: Title at Company | Date - Date
        pattern_a = r'(?P<title>[\w\s,./&-]+?)\s+(?:at|@)\s+(?P<company>[A-Z][\w\s&,.()-]+?)\s*[|\n]\s*(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\s*[-–—]\s*(?P<end>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current)'
        
        # Pattern B: Company Name \n Title \n Date - Date
        pattern_b = r'(?P<company>[A-Z][\w\s&,.()-]{2,50}?)\n\s*(?P<title>[\w\s,./&-]+?)\n\s*(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\s*[-–—]\s*(?P<end>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current)'
        
        # Pattern C: Title | Company | Date - Date
        pattern_c = r'(?P<title>[\w\s,./&-]+?)\s*[|\n]\s*(?P<company>[A-Z][\w\s&,.()-]+?)\s*[|\n]\s*(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\s*[-–—]\s*(?P<end>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current)'
        
        # Pattern D: Simple pattern without strict date format
        pattern_d = r'(?P<title>(?:Senior|Junior|Lead|Principal|Staff|Associate)?\s*(?:Software|Data|Machine Learning|Full Stack|Backend|Frontend|DevOps|Cloud)?\s*(?:Engineer|Developer|Analyst|Architect|Manager|Consultant|Designer|Lead))\s+(?:at|@|,|-)\s+(?P<company>[A-Z][\w\s&,.()-]{2,40})'
        
        patterns = [
            ('Pattern A', pattern_a),
            ('Pattern B', pattern_b),
            ('Pattern C', pattern_c),
            ('Pattern D', pattern_d)
        ]
        
        for pattern_name, pattern in patterns:
            matches = re.finditer(pattern, experience_section, re.IGNORECASE | re.MULTILINE)
            
            for match in matches:
                try:
                    designation = match.group('title').strip() if 'title' in match.groupdict() else None
                    company = match.group('company').strip() if 'company' in match.groupdict() else None
                    start_date = match.group('start').strip() if 'start' in match.groupdict() else None
                    end_date = match.group('end').strip() if 'end' in match.groupdict() else None
                    
                    if not designation or not company:
                        continue
                    
                    # Clean up
                    company = company.rstrip(',;.\n|')
                    company = re.sub(r'\s+', ' ', company)
                    designation = designation.rstrip(',;.\n|')
                    designation = re.sub(r'\s+', ' ', designation)
                    
                    # Validate
                    if len(company) < 2 or len(company) > 60 or len(designation) < 5:
                        continue
                    
                    # Deduplicate
                    key = (designation.lower().strip(), company.lower().strip())
                    if key in seen_entries:
                        continue
                    seen_entries.add(key)
                    
                    entry = {
                        'designation': designation,
                        'company': company
                    }
                    
                    # Add duration
                    if start_date and end_date:
                        entry['duration'] = f"{start_date} - {end_date}"
                    elif start_date:
                        entry['duration'] = f"{start_date} - Present"
                    
                    # Extract location from context
                    context_start = max(0, match.start())
                    context_end = min(len(experience_section), match.end() + 200)
                    context = experience_section[context_start:context_end]
                    
                    location_match = re.search(r'\n\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*(?:,\s*)?([A-Z][a-z]+)?', context)
                    if location_match:
                        location = location_match.group(0).strip()
                        if 3 < len(location) < 40:
                            entry['location'] = location
                    
                    # Extract key highlights (bullet points)
                    highlights = re.findall(r'[•●▪▸◦-]\s*(.+?)(?=\n|$)', context)
                    if highlights:
                        entry['highlights'] = [h.strip() for h in highlights[:5]]
                    
                    work_history.append(entry)
                    logger.info(f"Match with {pattern_name}: {designation} at {company}")
                    
                except Exception as e:
                    logger.warning(f"Error processing match: {e}")
                    continue
        
        logger.info(f"Extracted {len(work_history)} work experience entries")
        return work_history
    
    def parse(self, file_path: str) -> Dict:
        """
        Main parsing method
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Dictionary containing extracted information
        """
        try:
            logger.info(f"Starting parse of: {file_path}")
            
            # Extract text
            text = self.extract_text(file_path)
            
            if not text.strip():
                raise ValueError("No text could be extracted from the file")
            
            # Log extracted text for debugging (first 500 chars)
            logger.info(f"Extracted text preview: {text[:500]}")
            
            # Extract all information
            logger.info("Extracting information...")
            
            skills_data = self.extract_skills(text)
            education_data = self.extract_education(text)
            experience_data = self.extract_experience(text)
            work_history = self.extract_work_experience(text)
            
            result = {
                'name': self.extract_name(text),
                'email': self.extract_email(text),
                'phone_numbers': self.extract_phone(text),
                'skills': skills_data['all_skills'],
                'skills_categorized': skills_data['categorized'],
                'total_experience_years': experience_data['total_years'],
                'experience_details': experience_data,
                'education': education_data,
                'work_history': work_history,
                'parsed_at': datetime.now().isoformat(),
                'file_name': Path(file_path).name
            }
            
            logger.info("Parsing completed successfully")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing resume: {e}", exc_info=True)
            raise
    
    def parse_to_json(self, file_path: str, output_path: Optional[str] = None) -> str:
        """
        Parse resume and save as JSON
        
        Args:
            file_path: Path to resume file
            output_path: Optional path to save JSON (default: same name as input)
            
        Returns:
            JSON string of parsed data
        """
        result = self.parse(file_path)
        json_str = json.dumps(result, indent=2, ensure_ascii=False)
        
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(json_str)
            logger.info(f"Saved parsed data to: {output_path}")
        
        return json_str